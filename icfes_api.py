from flask import Flask, request, jsonify
from flask_cors import CORS
import google.generativeai as genai
import os
from dotenv import load_dotenv
import logging
from datetime import datetime
import json
import PyPDF2
from docx import Document
from io import BytesIO
import re
from werkzeug.utils import secure_filename
import tempfile                                                                                                                          
import matplotlib
matplotlib.use('Agg')  # usar backend no interactivo para servidores
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from io import BytesIO
import base64
import numpy as np
from charts import bar_chart_exam, pie_chart_exam

import sys
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass


# Cargar variables de entorno
load_dotenv()

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, origins=["*"], methods=["GET", "POST", "OPTIONS"])

# Configuración de API de IA (prioriza variable de entorno, luego fallback)
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") or "AIzaSyDBv91sukUg8BHVLde_1Jf5LjTNUyN4eKE"

if not GOOGLE_API_KEY:
    logger.error("❌ GOOGLE_API_KEY no encontrada")
    raise ValueError("Se requiere GOOGLE_API_KEY")

genai.configure(api_key=GOOGLE_API_KEY)

# ============================================================
# BASE DE DATOS SIMULADA
# ============================================================
users_db = {
    '1234567890': {
        'name': 'Usuario Demo',
        'password': 'demo123',
        'cedula': '1234567890',
        'created_at': str(datetime.now())
    }
}

# ============================================================
# CLASE EVALUADOR CON IA (GEMINI) - MEJORADO
# ============================================================
class GeminiEvaluator:
    def __init__(self):
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        self.default_model = "gemini-2.0-flash"
        self.custom_knowledge_base = {}  # Base de conocimiento personalizada
        self.trained_prompts = {}  # Prompts entrenados por competencia

    def _make_gemini_request(self, prompt, temperature=0.3, max_tokens=2000, max_retries=3):
        """Hace una petición a la API de Google Gemini con manejo robusto de errores"""
        for attempt in range(max_retries):
            try:
                logger.info(f"Enviando petición a Gemini (intento {attempt + 1}/{max_retries})")

                generation_config = genai.types.GenerationConfig(
                    temperature=temperature,
                    max_output_tokens=max_tokens
                )

                response = self.model.generate_content(
                    prompt,
                    generation_config=generation_config
                )

                if response and response.text:
                    return response.text

                raise Exception("Respuesta vacía de Gemini")

            except Exception as e:
                error_str = str(e).lower()

                if "rate limit" in error_str or "429" in error_str or "quota" in error_str:
                    if attempt < max_retries - 1:
                        wait_time = min((2 ** attempt) * 10, 60)
                        logger.warning(f"Rate limit alcanzado. Esperando {wait_time} segundos...")
                        import time
                        time.sleep(wait_time)
                        continue
                    else:
                        raise Exception("Rate limit excedido en Gemini")

                elif "invalid" in error_str or "400" in error_str:
                    raise Exception("Request inválido a Gemini")

                elif "api key" in error_str or "403" in error_str:
                    raise Exception("API Key de Gemini inválida")

                else:
                    logger.error(f"Error en petición a Gemini: {str(e)}")
                    if attempt == max_retries - 1:
                        raise

    def add_custom_knowledge(self, competencia, knowledge_text):
        """Añade conocimiento personalizado para una competencia específica"""
        if competencia not in self.custom_knowledge_base:
            self.custom_knowledge_base[competencia] = []
        self.custom_knowledge_base[competencia].append({
            'text': knowledge_text,
            'added_at': datetime.now().isoformat()
        })
        logger.info(f"Conocimiento personalizado añadido para competencia: {competencia}")

    def get_enhanced_prompt(self, competencia, base_prompt):
        """Mejora un prompt base con conocimiento personalizado"""
        enhanced_prompt = base_prompt

        if competencia in self.custom_knowledge_base:
            custom_knowledge = self.custom_knowledge_base[competencia]
            knowledge_text = "\n".join([item['text'] for item in custom_knowledge[-5:]])  # Últimos 5 elementos

            enhanced_prompt = f"""
CONTEXTO PERSONALIZADO PARA LA COMPETENCIA '{competencia}':
{knowledge_text}

INSTRUCCIONES ADICIONALES:
- Utiliza el contexto personalizado proporcionado arriba para generar preguntas más relevantes y precisas
- Asegúrate de que las preguntas reflejen el contenido específico enseñado
- Adapta las preguntas al nivel y enfoque del material proporcionado

{base_prompt}
"""

        return enhanced_prompt

    def analyze_document_content(self, text):
        """Analiza el documento completo y genera retroalimentación detallada"""
        prompt = f"""
Eres un experto pedagogo y analista educativo especializado en evaluaciones ICFES colombianas. Analiza el siguiente texto que contiene preguntas y respuestas de evaluación.

INSTRUCCIONES ESPECÍFICAS PARA ICFES:
1. Identifica TODAS las preguntas de opción múltiple siguiendo el formato ICFES
2. Para cada pregunta, genera retroalimentación educativa que incluya:
   - Explicación detallada de por qué la respuesta es correcta/incorrecta
   - Análisis de errores comunes en estudiantes colombianos
   - Sugerencias específicas de mejora basadas en el currículo ICFES
   - Conceptos clave del área de conocimiento ICFES

3. Evalúa el nivel de dominio según estándares ICFES:
   - Básico: Reconoce información básica
   - Satisfactorio: Aplica conocimientos en situaciones conocidas
   - Avanzado: Resuelve problemas complejos y transfiere conocimientos

FORMATO JSON obligatorio (responde con el JSON directamente, sin bloques de código markdown):
{{
    "analisis_completo": [
        {{
            "numero": 1,
            "pregunta": "Texto completo de la pregunta",
            "respuesta_elegida": "Opción seleccionada por el estudiante",
            "retroalimentacion": "Feedback educativo detallado con explicación del error y cómo mejorar",
            "respuesta_correcta": "La opción correcta con explicación detallada",
            "errores_comunes": "Errores típicos en estudiantes ICFES en esta área",
            "sugerencias_mejora": "Recomendaciones específicas basadas en currículo ICFES",
            "conceptos_clave": "Conceptos ICFES importantes relacionados",
            "nivel_icfes": "Básico/Satisfactorio/Avanzado"
        }}
    ],
    "resumen": {{
        "total_preguntas": 0,
        "documento_procesado": true,
        "observaciones": "Comentarios sobre el desempeño según estándares ICFES",
        "nivel_dominio_general": "Evaluación general: Básico/Satisfactorio/Avanzado",
        "recomendaciones_icfes": "Sugerencias específicas para mejorar en ICFES"
    }}
}}

IMPORTANTE: El JSON debe ser válido y parseable. No incluyas comas al final de listas u objetos.

TEXTO A ANALIZAR:
{text[:7000]}

Responde ÚNICAMENTE con el JSON válido.
"""
        try:
            response = self._make_gemini_request(prompt, temperature=0.4, max_tokens=6000)
            return response
        except Exception as e:
            logger.error(f"Error analizando documento con Gemini: {str(e)}")
            raise Exception(f"Error al analizar documento con Gemini: {str(e)}")

# Inicializar evaluador
evaluator = GeminiEvaluator()

# ============================================================
# UTILIDADES PARA PROCESAMIENTO DE ARCHIVOS
# ============================================================
def extract_text_from_pdf(file_content):
    """Extrae texto de un archivo PDF"""
    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""

        if len(pdf_reader.pages) == 0:
            raise Exception("El PDF no contiene páginas")

        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Página {page_num + 1} ---\n"
                    text += page_text + "\n"
            except Exception as e:
                logger.warning(f"Error extrayendo página {page_num + 1}: {str(e)}")
                continue

        if not text.strip():
            raise Exception("No se pudo extraer texto legible del PDF")

        return text.strip()
    except Exception as e:
        logger.error(f"Error procesando PDF: {str(e)}")
        raise Exception(f"Error al procesar PDF: {str(e)}")

def extract_text_from_docx(file_content):
    """Extrae texto de un archivo Word"""
    try:
        doc_file = BytesIO(file_content)
        doc = Document(doc_file)
        text = ""

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"

        if not text.strip():
            raise Exception("No se pudo extraer texto del documento Word")

        return text.strip()
    except Exception as e:
        logger.error(f"Error procesando Word: {str(e)}")
        raise Exception(f"Error al procesar documento Word: {str(e)}")

def clean_json_response(text):
    """Limpia y extrae JSON de la respuesta de IA"""
    try:
        text = text.strip()

        if text.startswith('```json'):
            text = text[7:].strip()
        elif text.startswith('```'):
            text = text[3:].strip()

        if text.endswith('```'):
            text = text[:-3].strip()

        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            json_text = json_match.group(0)
            json_text = re.sub(r',\s*([}\]])', r'\1', json_text)
            return json_text

        return text
    except Exception as e:
        logger.error(f"Error limpiando JSON: {str(e)}")
        return text

# ============================================================
# PARSEO ROBUSTO DE RESPUESTA JSON (PLAN DE FIGURAS)
# ============================================================
def _safe_parse_shapes_response(raw_text):
    """Intenta extraer una lista de 'shapes' desde diferentes formatos devueltos por IA.
    Devuelve una lista (posiblemente vacía) de shapes.
    """
    if not raw_text:
        return []

    # 1) Limpieza básica y extracción de bloque JSON
    cleaned = clean_json_response(raw_text)

    # 2) Intento directo
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, dict):
            return obj.get('shapes', []) or []
        if isinstance(obj, list):
            return obj  # A veces devuelve directamente la lista de shapes
    except Exception:
        pass

    # 3) Reparaciones simples: comillas simples y comas colgantes
    try:
        repaired = re.sub(r",\s*([}\]])", r"\1", cleaned)  # quitar comas finales
        # cambiar comillas simples por dobles solo cuando parezca JSON
        if '"' not in repaired and "'" in repaired:
            repaired = repaired.replace("'", '"')
        obj = json.loads(repaired)
        if isinstance(obj, dict):
            return obj.get('shapes', []) or []
        if isinstance(obj, list):
            return obj
    except Exception:
        pass

    # 4) Regex: extraer el array de shapes dentro de un objeto
    try:
        m = re.search(r'"shapes"\s*:\s*(\[.*?\])', cleaned, re.DOTALL)
        if m:
            arr_txt = m.group(1)
            arr_txt = re.sub(r",\s*]", "]", arr_txt)
            if '"' not in arr_txt and "'" in arr_txt:
                arr_txt = arr_txt.replace("'", '"')
            shapes = json.loads(arr_txt)
            if isinstance(shapes, list):
                return shapes
    except Exception:
        pass

    # 5) Fallback vacío
    return []

# ============================================================
# ENDPOINTS DE AUTENTICACIÓN
# ============================================================
@app.route('/register', methods=['POST'])
def register():
    """Registro de nuevos usuarios con cédula"""
    data = request.json
    cedula = data.get('cedula', '').strip()
    password = data.get('password', '')
    name = data.get('name', '').strip()
    
    if not cedula or not password or not name:
        return jsonify({'error': 'Todos los campos son requeridos'}), 400
    
    if not cedula.isdigit():
        return jsonify({'error': 'La cédula debe contener solo números'}), 400
    
    if cedula in users_db:
        return jsonify({'error': 'Esta cédula ya está registrada'}), 400
    
    if len(password) < 6:
        return jsonify({'error': 'La contraseña debe tener al menos 6 caracteres'}), 400
    
    users_db[cedula] = {
        'name': name,
        'password': password,
        'cedula': cedula,
        'created_at': str(datetime.now())
    }
    
    logger.info(f"Usuario registrado: {cedula} - {name}")
    
    return jsonify({
        'message': 'Usuario registrado exitosamente',
        'user': {
            'cedula': cedula,
            'name': name
        }
    }), 201

@app.route('/login', methods=['POST'])
def login():
    """Login de usuarios con cédula y rol"""
    data = request.json
    cedula = data.get('cedula', '').strip()
    password = data.get('password', '')
    role = data.get('role', 'estudiante').lower()
    
    if not cedula or not password:
        return jsonify({'error': 'Cédula y contraseña son requeridos'}), 400
    
    user = users_db.get(cedula)
    if not user:
        return jsonify({'error': 'Cédula no encontrada'}), 404
    
    if user['password'] != password:
        return jsonify({'error': 'Contraseña incorrecta'}), 401
    
    if role not in ['profesor', 'estudiante']:
        return jsonify({'error': 'Rol inválido. Debe ser profesor o estudiante'}), 400
    
    welcome_msg = 'Login exitoso - Profesor evaluador PDF' if role == 'profesor' else 'Login exitoso - Estudiante'
    
    logger.info(f"Login exitoso: {cedula} como {role}")
    
    return jsonify({
        'message': welcome_msg,
        'user': {
            'cedula': cedula,
            'name': user['name'],
            'role': role
        }
    }), 200

# ============================================================
# ENDPOINTS DE GENERACIÓN DE PREGUNTAS
# ============================================================
@app.route('/generate-question', methods=['POST'])
def generate_question():
    """Generar preguntas ICFES con diferentes niveles de dificultad usando conocimiento personalizado"""
    data = request.json
    competencia = data.get('competencia')
    num_questions = data.get('num_questions', 1)
    dificultad = data.get('dificultad', 'medio').lower()
    use_custom_knowledge = data.get('use_custom_knowledge', True)  # Por defecto usa conocimiento personalizado

    if not competencia:
        return jsonify({'error': 'Competencia es requerida'}), 400

    if num_questions < 1 or num_questions > 20:
        return jsonify({'error': 'El número de preguntas debe estar entre 1 y 20'}), 400

    if dificultad not in ['facil', 'medio', 'avanzado']:
        return jsonify({'error': 'Dificultad debe ser: facil, medio o avanzado'}), 400

    if dificultad == 'facil':
        nivel_texto = "básico/fácil"
        caracteristicas = """
   - Nivel de dificultad básico: conceptos fundamentales, preguntas directas
   - Requieren reconocimiento o recuerdo de información básica
   - Las opciones incorrectas son claramente erróneas
   - Evalúan conocimientos básicos y comprensión elemental"""
    elif dificultad == 'medio':
        nivel_texto = "intermedio"
        caracteristicas = """
   - Nivel de dificultad intermedio: aplicación de conceptos
   - Requieren análisis y comprensión de situaciones
   - Las opciones incorrectas deben ser plausibles pero incorrectas
   - Evalúan razonamiento lógico y aplicación de conocimientos"""
    else:
        nivel_texto = "avanzado/complejo"
        caracteristicas = """
   - Nivel de dificultad avanzado: síntesis y evaluación
   - Requieren análisis crítico, síntesis de información y resolución de problemas complejos
   - Las opciones incorrectas son muy plausibles y requieren discernimiento fino
   - Evalúan pensamiento crítico y aplicación avanzada de conceptos"""

    base_prompt = f"""Eres un experto en educación y diseño de evaluaciones ICFES colombianas.
Genera exactamente {num_questions} preguntas tipo ICFES de nivel {nivel_texto} para la competencia '{competencia}'.

INSTRUCCIONES CRÍTICAS PARA ICFES:
1. Cada pregunta debe seguir EXACTAMENTE este formato (sin variaciones):

Pregunta: [Escribe aquí la pregunta completa y clara]
a) [Primera opción]
b) [Segunda opción]
c) [Tercera opción]
d) [Cuarta opción]
Respuesta correcta: [letra de la opción correcta]
Explicación: [Explicación detallada de por qué es correcta esta respuesta y por qué las otras son incorrectas]


2. Características ICFES de las preguntas:
   - Deben ser relevantes para el currículo ICFES colombiano{caracteristicas}
   - Basadas en situaciones reales del contexto educativo colombiano
   - Redacción clara y sin ambigüedades según estándares ICFES

3. IMPORTANTE:
   - Genera EXACTAMENTE {num_questions} preguntas
   - Separa cada pregunta completa con DOS líneas en blanco
   - No agregues texto adicional fuera del formato
   - Todas las opciones deben ser distintas y plausibles
   - Las preguntas deben reflejar el nivel de complejidad ICFES correspondiente

Genera las preguntas ahora:"""

    # Usar conocimiento personalizado si está disponible y solicitado
    if use_custom_knowledge and competencia in evaluator.custom_knowledge_base:
        prompt = evaluator.get_enhanced_prompt(competencia, base_prompt)
        logger.info(f"Usando conocimiento personalizado para competencia: {competencia}")
    else:
        prompt = base_prompt
        logger.info(f"Generando preguntas sin conocimiento personalizado para: {competencia}")

    try:
        questions_text = evaluator._make_gemini_request(
            prompt,
            temperature=0.7,
            max_tokens=4096
        )

        if not questions_text or len(questions_text) < 50:
            return jsonify({'error': 'No se pudieron generar preguntas válidas'}), 500

        logger.info(f"Preguntas generadas: {competencia} - {dificultad} - {num_questions} - Custom Knowledge: {use_custom_knowledge}")

        return jsonify({
            'questions': questions_text,
            'custom_knowledge_used': use_custom_knowledge and competencia in evaluator.custom_knowledge_base,
            'competencia': competencia,
            'dificultad': dificultad,
            'num_questions': num_questions
        }), 200

    except Exception as e:
        logger.error(f"Error generando preguntas: {str(e)}")
        return jsonify({'error': f'Error al generar preguntas: {str(e)}'}), 500

# ============================================================
# ENDPOINTS DE RETROALIMENTACIÓN
# ============================================================
@app.route('/get-feedback', methods=['POST'])
def get_feedback():
    """Obtener retroalimentación detallada de la respuesta del estudiante"""
    data = request.json
    question = data.get('question')
    answer = data.get('answer')
    
    if not question or not answer:
        return jsonify({'error': 'Pregunta y respuesta son requeridas'}), 400

    prompt = f"""Eres un tutor experto y pedagógico. Analiza la siguiente pregunta ICFES y la respuesta del estudiante.

PREGUNTA:
{question}

RESPUESTA DEL ESTUDIANTE: {answer}

Proporciona retroalimentación siguiendo EXACTAMENTE este formato:

**Estado de la respuesta:** [Indica si es Correcta o Incorrecta]

**Análisis de tu respuesta:**
[Explica por qué eligieron esa opción y qué razonamiento pudieron haber seguido]

**Respuesta correcta:**
[Indica cuál es la respuesta correcta con su letra y el texto completo de la opción]

**Explicación detallada:**
[Explica claramente por qué esa es la respuesta correcta, haciendo referencia al contenido de la pregunta]

**En qué fallaste:**
[Si la respuesta fue incorrecta, explica específicamente el error conceptual o de razonamiento. Si fue correcta, explica qué hizo bien el estudiante]

**Cómo mejorar:**
[Proporciona consejos específicos y prácticos para abordar este tipo de preguntas en el futuro. Incluye estrategias de estudio y comprensión]

**Conceptos clave a repasar:**
[Lista 2-3 conceptos específicos que el estudiante debe estudiar o reforzar]

Sé claro, constructivo y motivador en tu retroalimentación."""

    try:
        feedback_text = evaluator._make_gemini_request(
            prompt,
            temperature=0.6,
            max_tokens=2048
        )
        
        if not feedback_text:
            return jsonify({'error': 'No se pudo generar retroalimentación'}), 500
        
        logger.info("Retroalimentación generada exitosamente")
        
        return jsonify({'feedback': feedback_text}), 200
        
    except Exception as e:
        logger.error(f"Error generando retroalimentación: {str(e)}")
        return jsonify({'error': f'Error al generar retroalimentación: {str(e)}'}), 500

# ============================================================
# ENDPOINTS DE ANÁLISIS DE DOCUMENTOS
# ============================================================
@app.route('/analyze-document', methods=['POST'])
def analyze_document():
    """Analiza un documento completo (PDF o Word) con validación robusta"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No se proporcionó ningún archivo"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No se seleccionó ningún archivo"}), 400

        allowed_types = ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']
        if file.content_type not in allowed_types:
            return jsonify({
                "error": "Tipo de archivo no soportado",
                "supported_formats": ["PDF", "Word (.docx, .doc)"],
                "received_type": file.content_type
            }), 400

        file_content = file.read()

        if len(file_content) > 10 * 1024 * 1024:
            return jsonify({"error": "El archivo es demasiado grande. Máximo 10MB."}), 400

        if len(file_content) < 100:
            return jsonify({"error": "El archivo está vacío o es demasiado pequeño"}), 400

        logger.info(f"Procesando archivo: {file.filename}, tamaño: {len(file_content)} bytes")

        filename = file.filename.lower()
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
        elif filename.endswith(('.docx', '.doc')):
            text = extract_text_from_docx(file_content)
        else:
            return jsonify({"error": "Extensión de archivo no reconocida"}), 400

        if len(text.strip()) < 50:
            return jsonify({"error": "El archivo no contiene suficiente texto para procesar"}), 400

        logger.info(f"Texto extraído: {len(text)} caracteres")

        try:
            ai_response = evaluator.analyze_document_content(text)
            clean_response = clean_json_response(ai_response)

            logger.info(f"Respuesta IA (primeros 300 chars): {clean_response[:300]}...")

            analysis_data = json.loads(clean_response)
            analysis_results = analysis_data.get('analisis_completo', [])
            resumen = analysis_data.get('resumen', {})

            if not analysis_results:
                return jsonify({
                    "error": "No se encontraron preguntas para analizar en el documento",
                    "suggestion": "Verifica que el documento contenga preguntas de evaluación claras"
                }), 404

            logger.info(f"Análisis completado: {len(analysis_results)} elementos")

            return jsonify({
                "success": True,
                "analysis": analysis_results,
                "summary": resumen,
                "metadata": {
                    "file_processed": file.filename,
                    "file_size": len(file_content),
                    "text_length": len(text),
                    "total_items": len(analysis_results),
                    "processed_at": datetime.now().isoformat()
                }
            })

        except json.JSONDecodeError as e:
            logger.error(f"Error JSON: {str(e)}")
            logger.error(f"Respuesta problemática: {clean_response[:500]}...")
            return jsonify({
                "error": "Error procesando la respuesta de IA",
                "details": f"Respuesta inválida: {str(e)}",
                "suggestion": "Intenta con un documento más claro o reformateado"
            }), 500

    except Exception as e:
        logger.error(f"Error general en analyze_document: {str(e)}")
        return jsonify({
            "error": "Error interno del servidor",
            "details": str(e),
            "type": type(e).__name__
        }), 500

# ============================================================
# ENDPOINTS DE CONOCIMIENTO PERSONALIZADO
# ============================================================
@app.route('/add-custom-knowledge', methods=['POST'])
def add_custom_knowledge():
    """Añadir conocimiento personalizado para mejorar las preguntas ICFES"""
    data = request.json
    competencia = data.get('competencia', '').strip()
    knowledge_text = data.get('knowledge_text', '').strip()

    if not competencia:
        return jsonify({'error': 'Competencia es requerida'}), 400

    if not knowledge_text:
        return jsonify({'error': 'Texto de conocimiento es requerido'}), 400

    if len(knowledge_text) < 50:
        return jsonify({'error': 'El texto debe tener al menos 50 caracteres'}), 400

    try:
        evaluator.add_custom_knowledge(competencia, knowledge_text)

        # Contar elementos en la base de conocimiento
        knowledge_count = len(evaluator.custom_knowledge_base.get(competencia, []))

        logger.info(f"Conocimiento personalizado añadido: {competencia} - Total elementos: {knowledge_count}")

        return jsonify({
            'success': True,
            'message': f'Conocimiento personalizado añadido para {competencia}',
            'competencia': competencia,
            'knowledge_count': knowledge_count,
            'added_at': datetime.now().isoformat()
        }), 201

    except Exception as e:
        logger.error(f"Error añadiendo conocimiento personalizado: {str(e)}")
        return jsonify({'error': f'Error al añadir conocimiento: {str(e)}'}), 500

@app.route('/get-custom-knowledge', methods=['GET'])
def get_custom_knowledge():
    """Obtener el conocimiento personalizado disponible"""
    competencia = request.args.get('competencia')

    if competencia:
        knowledge = evaluator.custom_knowledge_base.get(competencia, [])
        return jsonify({
            'competencia': competencia,
            'knowledge_items': knowledge,
            'total_items': len(knowledge)
        }), 200
    else:
        # Retornar resumen de todas las competencias
        summary = {}
        for comp, items in evaluator.custom_knowledge_base.items():
            summary[comp] = {
                'total_items': len(items),
                'last_updated': items[-1]['added_at'] if items else None
            }

        return jsonify({
            'custom_knowledge_summary': summary,
            'total_competencias': len(summary)
        }), 200

# ============================================================
# ENDPOINTS DE GUARDADO DE MODELOS
# ============================================================
@app.route('/save-model', methods=['POST'])
def save_model():
    """Guardar modelo entrenado en archivo JSON"""
    try:
        data = request.json
        model_data = data.get('model_data')
        file_name = data.get('file_name', 'modelo_entrenado')
        timestamp = data.get('timestamp', datetime.now().isoformat())

        if not model_data:
            return jsonify({'error': 'Datos del modelo son requeridos'}), 400

        # Crear nombre de archivo seguro
        safe_filename = secure_filename(file_name)
        if not safe_filename.endswith('.json'):
            safe_filename += '.json'

        # Crear directorio models si no existe
        models_dir = 'models'
        if not os.path.exists(models_dir):
            os.makedirs(models_dir)

        # Guardar modelo con timestamp
        model_filename = f"{models_dir}/modelo_{int(datetime.now().timestamp())}_{safe_filename}"
        model_info = {
            'model_data': model_data,
            'original_file': file_name,
            'saved_at': timestamp,
            'server_timestamp': datetime.now().isoformat()
        }

        with open(model_filename, 'w', encoding='utf-8') as f:
            json.dump(model_info, f, ensure_ascii=False, indent=2)

        logger.info(f"Modelo guardado: {model_filename}")

        return jsonify({
            'success': True,
            'message': 'Modelo entrenado guardado exitosamente',
            'filename': model_filename,
            'saved_at': datetime.now().isoformat()
        }), 200

    except Exception as e:
        logger.error(f"Error guardando modelo: {str(e)}")
        return jsonify({'error': f'Error al guardar el modelo: {str(e)}'}), 500

@app.route('/generate-visual', methods=['POST'])
def generate_visual():
    """Genera gráficos visuales (bar/pie) para análisis de preguntas ICFES usando Gemini para datos"""
    data = request.json
    chart_type = data.get('chart_type', 'bar').lower()
    analysis_data = data.get('analysis_data', '')
    custom_labels = data.get('labels', [])
    custom_values = data.get('values', [])
    title = data.get('title', 'Análisis de Respuestas ICFES')
    xlabel = data.get('xlabel', 'Categorías')
    ylabel = data.get('ylabel', 'Número de Preguntas')
    # Opciones de estilo y formato
    paleta_bw = bool(data.get('paleta_bw', True))
    donut = bool(data.get('donut', True))
    ordenar_desc = bool(data.get('ordenar_desc', True))
    force_percentage = data.get('porcentaje')  # None (auto) o bool

    if chart_type not in ['bar', 'pie']:
        return jsonify({'error': 'Tipo de gráfico no soportado. Use "bar" o "pie"'}), 400

    try:
        # Si no se proporcionan datos personalizados, usar Gemini para generarlos
        if not custom_labels or not custom_values:
            if not analysis_data:
                return jsonify({'error': 'Se requiere analysis_data o labels/values personalizados'}), 400

            # Prompt para Gemini para generar datos de visualización
            prompt = f"""
Eres un analista educativo experto. Analiza los siguientes datos de evaluación ICFES y genera datos para un gráfico de {chart_type}.

DATOS DE ANÁLISIS:
{analysis_data}

INSTRUCCIONES:
1. Identifica categorías relevantes para el gráfico (ej. tipos de errores, competencias, niveles de dificultad)
2. Genera valores numéricos realistas basados en el análisis
3. Proporciona exactamente 4-6 categorías con valores porcentuales o absolutos

FORMATO JSON obligatorio (responde con el JSON directamente, sin bloques de código markdown):
{{
    "labels": ["Categoría 1", "Categoría 2", "Categoría 3", "Categoría 4"],
    "values": [25, 30, 20, 25],
    "title": "Título sugerido para el gráfico",
    "xlabel": "Etiqueta eje X",
    "ylabel": "Etiqueta eje Y"
}}

IMPORTANTE: El JSON debe ser válido y parseable. Los valores deben ser números enteros.
"""

            ai_response = evaluator._make_gemini_request(prompt, temperature=0.3, max_tokens=1000)
            clean_response = clean_json_response(ai_response)

            try:
                chart_data = json.loads(clean_response)
                labels = chart_data.get('labels', [])
                values = chart_data.get('values', [])
                if chart_data.get('title'):
                    title = chart_data['title']
                if chart_data.get('xlabel'):
                    xlabel = chart_data['xlabel']
                if chart_data.get('ylabel'):
                    ylabel = chart_data['ylabel']
            except json.JSONDecodeError:
                return jsonify({'error': 'Error procesando respuesta de IA para datos del gráfico'}), 500
        else:
            labels = custom_labels
            values = custom_values

        if not labels or not values or len(labels) != len(values):
            return jsonify({'error': 'No se pudieron generar labels y values válidos'}), 400

        # Generar el gráfico con estilo editorial desde charts.py
        # Detección automática de porcentaje, a menos que el usuario lo fuerce
        if isinstance(force_percentage, bool):
            is_percentage = force_percentage
        else:
            is_percentage = False
            try:
                total_vals = sum(values)
                max_val = max(values)
                is_percentage = abs(total_vals - 100) < 1e-6 or max_val <= 1
            except Exception:
                is_percentage = False

        if chart_type == 'bar':
            fig, ax = bar_chart_exam(
                categorias=labels,
                valores=values,
                titulo=title,
                xlabel=xlabel,
                ylabel=ylabel if not is_percentage else 'Porcentaje',
                ordenar_desc=ordenar_desc,
                mostrar_valores=True,
                valores_como_porcentaje=is_percentage,
                paleta_bw=paleta_bw,
                salida=None,
            )
        else:
            fig, ax = pie_chart_exam(
                labels=labels,
                valores=values,
                titulo=title,
                donut=donut,
                paleta_bw=paleta_bw,
                salida=None,
            )

        img_buffer = BytesIO()
        fig.savefig(img_buffer, format='png', bbox_inches='tight', dpi=300)
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.read()).decode('utf-8')
        plt.close(fig)

        logger.info(f"Gráfico generado con Gemini: {chart_type} con {len(labels)} elementos")

        return jsonify({
            'success': True,
            'image': f'data:image/png;base64,{img_base64}',
            'chart_type': chart_type,
            'title': title,
            'labels': labels,
            'values': values,
            'generated_by_ai': True if not custom_labels else False
        }), 200

    except Exception as e:
        logger.error(f"Error generando gráfico: {str(e)}")
        plt.close('all')
        return jsonify({'error': f'Error al generar el gráfico: {str(e)}'}), 500

# ============================================================
# ENDPOINT: VISUAL POR COMPETENCIA (AGREGADO)
# ============================================================
@app.route('/generate-visual-by-competencia', methods=['POST'])
def generate_visual_by_competencia():
    """Agrupa elementos por 'competencia' y genera un gráfico de barras editorial.

    Entrada JSON esperada:
    {
      "items": [{"competencia": "Matemáticas", "valor": 12}, ...],
      "title": "Distribución por competencia",
      "paleta_bw": true,
      "ordenar_desc": true,
      "porcentaje": null | true | false
    }
    Si "valor" no se provee, se cuenta como 1 por ítem.
    "porcentaje" controla si mostrar porcentaje (true/false) o autodetección (null/omitido).
    "paleta_bw" aplica estilo en escala de grises.
    "ordenar_desc" controla orden.
    "xlabel" y "ylabel" opcionales.
    "ancho_mm" y "alto_mm" opcionales.
    "salida" se ignora y siempre se retorna PNG base64.
    """
    try:
        data = request.json or {}
        items = data.get('items', [])
        title = data.get('title', 'Distribución por competencia')
        xlabel = data.get('xlabel', 'Competencia')
        ylabel = data.get('ylabel', 'Frecuencia')
        paleta_bw = bool(data.get('paleta_bw', True))
        ordenar_desc = bool(data.get('ordenar_desc', True))
        force_percentage = data.get('porcentaje')  # None auto, o bool

        if not isinstance(items, list) or len(items) == 0:
            return jsonify({'error': 'Se requiere un arreglo "items" con objetos que tengan "competencia" y opcionalmente "valor"'}), 400

        # Agregación
        agg = {}
        for it in items:
            comp = str((it.get('competencia') or '')).strip()
            if not comp:
                # Saltar entradas sin competencia
                continue
            v = it.get('valor')
            try:
                val = float(v) if v is not None else 1.0
            except Exception:
                val = 1.0
            agg[comp] = agg.get(comp, 0.0) + val

        if not agg:
            return jsonify({'error': 'Tras la agregación no hay datos válidos. Verifica el campo "competencia" en los items.'}), 400

        labels = list(agg.keys())
        values = list(agg.values())

        # Porcentaje auto o forzado
        if isinstance(force_percentage, bool):
            is_percentage = force_percentage
        else:
            try:
                total_vals = sum(values)
                max_val = max(values)
                is_percentage = abs(total_vals - 100) < 1e-6 or max_val <= 1
            except Exception:
                is_percentage = False

        # Generar gráfico
        fig, ax = bar_chart_exam(
            categorias=labels,
            valores=values,
            titulo=title,
            xlabel=xlabel,
            ylabel=ylabel if not is_percentage else 'Porcentaje',
            ordenar_desc=ordenar_desc,
            mostrar_valores=True,
            valores_como_porcentaje=is_percentage,
            paleta_bw=paleta_bw,
            salida=None,
        )

        buf = BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', dpi=300)
        buf.seek(0)
        img_b64 = base64.b64encode(buf.read()).decode('utf-8')
        plt.close(fig)

        return jsonify({
            'success': True,
            'image': f'data:image/png;base64,{img_b64}',
            'labels': labels,
            'values': values,
            'title': title,
            'grouped_by': 'competencia'
        }), 200

    except Exception as e:
        logger.error(f"Error en generate-visual-by-competencia: {str(e)}")
        plt.close('all')
        return jsonify({'error': f'Error al generar el gráfico por competencia: {str(e)}'}), 500

# ============================================================
# ENDPOINT DE FIGURAS GEOMÉTRICAS POR PREGUNTA
# ============================================================
def _draw_geometry(ax, shapes):
    """Dibuja figuras en un eje Matplotlib a partir de una lista de shapes JSON."""
    # Config base
    ax.set_aspect('equal', adjustable='datalim')
    ax.grid(True, linestyle='--', alpha=0.3)

    # Colores y estilo
    color_main = '#667eea'
    color_aux = '#10b981'

    # Utilidades
    def as_point(p):
        return (float(p[0]), float(p[1]))

    for shape in shapes:
        try:
            stype = (shape.get('type') or '').lower()
            if stype == 'triangle':
                pts = [as_point(p) for p in shape.get('points', [])][:3]
                if len(pts) == 3:
                    poly = patches.Polygon(pts, closed=True, fill=False, edgecolor=color_main, linewidth=2)
                    ax.add_patch(poly)
            elif stype == 'polygon':
                pts = [as_point(p) for p in shape.get('points', [])]
                if len(pts) >= 3:
                    poly = patches.Polygon(pts, closed=True, fill=False, edgecolor=color_main, linewidth=2)
                    ax.add_patch(poly)
            elif stype == 'rectangle':
                x, y = as_point(shape.get('xy', [0, 0]))
                w = float(shape.get('width', 1))
                h = float(shape.get('height', 1))
                rect = patches.Rectangle((x, y), w, h, fill=False, edgecolor=color_main, linewidth=2)
                ax.add_patch(rect)
            elif stype == 'circle':
                cx, cy = as_point(shape.get('center', [0, 0]))
                r = float(shape.get('radius', 1))
                circ = patches.Circle((cx, cy), r, fill=False, edgecolor=color_main, linewidth=2)
                ax.add_patch(circ)
            elif stype == 'arc':
                cx, cy = as_point(shape.get('center', [0, 0]))
                r = float(shape.get('radius', 1))
                t1 = float(shape.get('theta1', 0))
                t2 = float(shape.get('theta2', 90))
                arc = patches.Arc((cx, cy), 2*r, 2*r, angle=0, theta1=t1, theta2=t2, color=color_main, linewidth=2)
                ax.add_patch(arc)
            elif stype == 'angle':
                atx, aty = as_point(shape.get('at', [0, 0]))
                radius = float(shape.get('radius', 1))
                start = float(shape.get('start_deg', 0))
                end = float(shape.get('end_deg', 60))
                # arco del ángulo
                arc = patches.Arc((atx, aty), 2*radius, 2*radius, angle=0, theta1=start, theta2=end, color=color_aux, linewidth=2)
                ax.add_patch(arc)
                # rayos del ángulo
                sx = atx + radius * np.cos(np.deg2rad(start))
                sy = aty + radius * np.sin(np.deg2rad(start))
                ex = atx + radius * np.cos(np.deg2rad(end))
                ey = aty + radius * np.sin(np.deg2rad(end))
                ax.plot([atx, sx], [aty, sy], color=color_aux, linewidth=2)
                ax.plot([atx, ex], [aty, ey], color=color_aux, linewidth=2)
            elif stype == 'segment':
                p1 = as_point(shape.get('p1', [0, 0]))
                p2 = as_point(shape.get('p2', [1, 0]))
                ax.plot([p1[0], p2[0]], [p1[1], p2[1]], color=color_main, linewidth=2)
            elif stype == 'point':
                px, py = as_point(shape.get('at', [0, 0]))
                label = shape.get('label')
                ax.scatter([px], [py], color=color_aux, s=30)
                if label:
                    ax.text(px + 0.05, py + 0.05, str(label), color=color_aux, fontsize=9)
            elif stype == 'axis' or stype == 'axes':
                # Ejes cartesianos básicos
                rng = float(shape.get('range', 5))
                ax.axhline(0, color='#999', linewidth=1)
                ax.axvline(0, color='#999', linewidth=1)
                ax.set_xlim(-rng, rng)
                ax.set_ylim(-rng, rng)
                # labels opcionales
                if shape.get('xlabel'):
                    ax.set_xlabel(str(shape['xlabel']))
                if shape.get('ylabel'):
                    ax.set_ylabel(str(shape['ylabel']))
            elif stype == 'function':
                # Plot de función y = f(x), ej: sin(x), x**2
                expr = str(shape.get('expr', ''))
                x_range = shape.get('x_range', [-10, 10])
                samples = int(shape.get('samples', 500))
                color = shape.get('color', color_main)
                try:
                    xs = np.linspace(float(x_range[0]), float(x_range[1]), samples)
                    # Ambiente seguro para eval
                    safe = { 'np': np, 'sin': np.sin, 'cos': np.cos, 'tan': np.tan, 'pi': np.pi, 'exp': np.exp, 'sqrt': np.sqrt }
                    ys = eval(expr, {"__builtins__": {}}, safe | {'x': xs})
                    ax.plot(xs, ys, color=color, linewidth=2)
                except Exception:
                    continue
            elif stype == 'wave':
                # Onda senoidal: A*sin(2*pi*f*x + phase)
                A = float(shape.get('amplitude', 1))
                f = float(shape.get('frequency', 1))
                phase = float(shape.get('phase', 0))
                x_range = shape.get('x_range', [0, 2*np.pi])
                samples = int(shape.get('samples', 800))
                color = shape.get('color', color_main)
                xs = np.linspace(float(x_range[0]), float(x_range[1]), samples)
                ys = A * np.sin(2*np.pi*f*xs + phase)
                ax.plot(xs, ys, color=color, linewidth=2)
            # Se pueden añadir más tipos: angle, arc, polygon, altitude, median, etc.
        except Exception as _:
            continue

    ax.autoscale(enable=True, tight=True)


@app.route('/generate-geometry-visual', methods=['POST'])
def generate_geometry_visual():
    """Genera una figura geométrica basada en el texto de la pregunta usando un plan JSON devuelto por IA."""
    data = request.json or {}
    question_text = data.get('question_text', '').strip()
    title = data.get('title', 'Figura Geométrica')
    competencia = (data.get('competencia') or '').strip()
    tema = (data.get('tema') or '').strip()

    if not question_text:
        return jsonify({'error': 'question_text es requerido'}), 400

    try:
        # Pedir a Gemini un plan de dibujo en JSON
        prompt = f"""
Eres un asistente de geometría. A partir del enunciado de una pregunta, devuelve un PLAN DE DIBUJO en JSON para visualizar la figura.

Reglas:
- No expliques nada; responde SOLO JSON válido y parseable.
- Usa tipos: triangle(points:[[x,y],[x,y],[x,y]]), polygon(points:[[x,y],...]), rectangle(xy:[x,y], width, height), circle(center:[x,y], radius), arc(center:[x,y], radius, theta1, theta2), angle(at:[x,y], radius, start_deg, end_deg), segment(p1:[x,y], p2:[x,y]), point(at:[x,y], label), axis(range, xlabel?, ylabel?), function(expr:"sin(x)", x_range:[-10,10], samples?, color?), wave(amplitude, frequency, phase?, x_range?, samples?, color?).
- Incluye ejes si ayuda a la comprensión.
- Ajusta medidas razonables si el enunciado no da números exactos.
- Usa preferentemente coordenadas ENTERAS y mantén todas las coordenadas y radios en el rango [-10, 10].
- Si el enunciado indica letras de puntos (A, B, C, etc.), agrega labels en los puntos.
- Si se menciona ángulo, lados, alturas o medianas, incluye los segmentos relevantes.

Formato JSON:
{{
  "shapes": [
    {{"type":"triangle","points":[[0,0],[4,0],[2,3]]}},
    {{"type":"point","at":[2,0],"label":"A"}},
    {{"type":"axis","range":6, "xlabel":"x", "ylabel":"y"}},
    {{"type":"function","expr":"sin(x)","x_range":[-3.14,3.14],"samples":400}}
  ]
}}

Enunciado:
{question_text}

Contexto (opcional):
Competencia: {competencia}
Tema: {tema}
"""

        ai_resp = evaluator._make_gemini_request(prompt, temperature=0.2, max_tokens=800)
        shapes = _safe_parse_shapes_response(ai_resp)

        # Si no se pudo interpretar, reintentar una vez con prompt ultra estricto
        if not shapes:
            strict_prompt = f"""
Responde EXCLUSIVAMENTE con un JSON válido. Estructura EXACTA: {{"shapes": [ .. ]}} sin comentarios, sin markdown, sin texto extra.
Usa los mismos tipos permitidos (triangle, polygon, rectangle, circle, arc, angle, segment, point, axis, function, wave).
Ajusta coordenadas a enteros dentro de [-10,10].
Enunciado:
{question_text}
"""
            ai_resp2 = evaluator._make_gemini_request(strict_prompt, temperature=0.0, max_tokens=500)
            shapes = _safe_parse_shapes_response(ai_resp2)

        # Renderizar figura
        fig, ax = plt.subplots(figsize=(6, 6))
        if shapes:
            _draw_geometry(ax, shapes)
        else:
            # Fallback seguro: solo ejes para no dejar vacío
            _draw_geometry(ax, [{"type": "axis", "range": 5, "xlabel": "x", "ylabel": "y"}])
        ax.set_title(title)

        buf = BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=120)
        buf.seek(0)
        img_b64 = base64.b64encode(buf.read()).decode('utf-8')
        plt.close(fig)

        logger.info("Figura geométrica generada correctamente")
        return jsonify({
            'success': True,
            'image': f'data:image/png;base64,{img_b64}',
            'shapes_count': len(shapes)
        })

    except json.JSONDecodeError:
        return jsonify({'error': 'No se pudo interpretar el plan JSON devuelto por IA'}), 500
    except Exception as e:
        logger.error(f"Error generando figura geométrica: {str(e)}")
        plt.close('all')
        return jsonify({'error': f'Error generando figura: {str(e)}'}), 500

# ============================================================
# ENDPOINTS ADMINISTRATIVOS
# ============================================================
@app.route('/users', methods=['GET'])
def get_users():
    """Endpoint para ver usuarios registrados (solo para desarrollo)"""
    users_list = []
    for cedula, user_data in users_db.items():
        users_list.append({
            'cedula': cedula,
            'name': user_data['name'],
            'created_at': user_data.get('created_at', 'N/A')
        })
    return jsonify({'users': users_list, 'total': len(users_list)}), 200

@app.route('/health', methods=['GET'])
def health_check():
    """Endpoint para verificar el estado del sistema"""
    try:
        test_response = evaluator._make_gemini_request(
            "Responde solo 'OK' para confirmar conexión.",
            max_tokens=10
        )
        ai_status = "OK" in test_response.upper()

        return jsonify({
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "server": "Flask running",
            "ai_api": "connected" if ai_status else "error",
            "ai_provider": "Gemini",
            "model": evaluator.default_model,
            "users_registered": len(users_db),
            "version": "ICFES Pro Backend Unificado v1.0"
        })
    except Exception as e:
        return jsonify({
            "status": "partial",
            "server": "Flask running",
            "ai_api": "error",
            "ai_provider": "Gemini",
            "error": str(e),
            "users_registered": len(users_db)
        }), 200

# ============================================================
# INICIO DE LA APLICACIÓN
# ============================================================
if __name__ == '__main__':
    print("=" * 60)
    print("🚀 ICFES PRO - BACKEND UNIFICADO")
    print("=" * 60)
    print(f"🔑 Google Gemini API: {'✅ Configurada' if GOOGLE_API_KEY else '❌ No encontrada'}")
    print(f"👥 Usuario Demo: Cédula: 1234567890 | Password: demo123")
    print("🌐 Servidor: http://127.0.0.1:5000")
    print("=" * 60)
    print("\n📋 ENDPOINTS DISPONIBLES:")
    print("\n🔐 Autenticación:")
    print("   - POST /register              - Registro de usuarios")
    print("   - POST /login                 - Login con roles")
    print("\n📝 Generación de Preguntas:")
    print("   - POST /generate-question     - Generar preguntas ICFES")
    print("   - POST /get-feedback          - Retroalimentación individual")
    print("\n📄 Análisis de Documentos:")
    print("   - POST /analyze-document      - Análisis PDF/Word completo")
    print("\n📊 Visualización de Datos:")
    print("   - POST /generate-visual       - Generar gráficos (bar/pie)")
    print("\n💾 Modelos de IA:")
    print("   - POST /save-model            - Guardar modelo entrenado")
    print("\n⚙️  Administrativo:")
    print("   - GET  /users                 - Lista de usuarios")
    print("   - GET  /health                - Estado del sistema")
    print("=" * 60)
    
    app.run(debug=True, host='127.0.0.1', port=5000)